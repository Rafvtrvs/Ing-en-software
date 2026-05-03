#!/usr/bin/env python3
"""Genera correcciones-documento-0-rubrica-detallado.docx desde el .md del repo."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_bold_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def add_justified_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_bold_runs(p, text)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def is_table_separator(s: str) -> bool:
    t = s.strip().strip("|")
    parts = [p.strip() for p in t.split("|")]
    if not parts:
        return False
    for p in parts:
        if p == "":
            continue
        if not re.match(r"^[-:]+$", p):
            return False
    return len(parts) >= 2


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    md_path = root / "correcciones-documento-0-rubrica-detallado.md"
    out_path = root / "correcciones-documento-0-rubrica-detallado.docx"

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        if is_table_row(line):
            rows_raw = []
            while i < len(lines) and is_table_row(lines[i]):
                rows_raw.append(lines[i])
                i += 1
            data_rows = []
            for r in rows_raw:
                if is_table_separator(r):
                    continue
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                data_rows.append(cells)
            if data_rows:
                ncols = max(len(r) for r in data_rows)
                for r in data_rows:
                    while len(r) < ncols:
                        r.append("")
                table = doc.add_table(rows=len(data_rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row_cells in enumerate(data_rows):
                    for ci, val in enumerate(row_cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        if ri == 0:
                            add_bold_runs(p, val)
                            set_cell_shading(cell, "D9E2F3")
                        else:
                            p.add_run(val)
                doc.add_paragraph()
            continue

        if stripped.startswith("- "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_bold_runs(p, content)
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num and not stripped.startswith("##"):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_bold_runs(p, m_num.group(2))
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") == 2:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(stripped[1:-1]).italic = True
            i += 1
            continue

        add_justified_paragraph(doc, stripped)
        i += 1

    doc.save(out_path)
    print(f"Escrito: {out_path}")


if __name__ == "__main__":
    main()
